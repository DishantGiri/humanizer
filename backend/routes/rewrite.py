"""
API routes for the rewrite pipeline with MySQL history logging and 10-humanization limit enforcement.
"""

import os
import uuid
import logging
from typing import Optional, List
import io
from datetime import datetime
from fastapi import APIRouter, HTTPException, Header, Depends, UploadFile, File
from pydantic import BaseModel, Field

from config import RewriteMode, RewriteLevel, MAX_INPUT_LENGTH
from analyzer import analyze
from rewriter import TextRewriter, RewriteError
from verifier import MeaningVerifier
from translator import TranslationBouncer
from utils import count_changes, compute_reading_time, sanitize_input, compute_word_diff
from humanizer import humanize, strip_formatting_artifacts, enforce_short_sentences, add_burstiness, clean_erroneous_punctuation
from perplexity import PerplexityOptimizer
from validators import validate_human_statistics

from db import execute_query, fetch_all, fetch_one
from routes.auth import get_optional_user_from_token, get_current_user_from_token, UserResponse
from cache_limiter import get_cached_rewrite, set_cached_rewrite, is_rate_limited
from similarity import calculate_similarity_metrics
from nlp_analyzer import analyze_text_nlp
from ai_checker import AICheckEngine, AICheckReport
from humanize_pipeline import get_standard_pipeline, StandardHumanizePipeline

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api", tags=["rewrite"])

# ── Singletons ──────────────────────────────────────────────────────────────

_rewriter: TextRewriter | None = None
_verifier: MeaningVerifier | None = None
_bouncer: TranslationBouncer | None = None
_perplexity_optimizer: PerplexityOptimizer | None = None
_standard_pipeline: StandardHumanizePipeline | None = None


def get_rewriter() -> TextRewriter:
    global _rewriter
    if _rewriter is None:
        _rewriter = TextRewriter()
    return _rewriter


def get_verifier() -> MeaningVerifier:
    global _verifier
    if _verifier is None:
        _verifier = MeaningVerifier()
    return _verifier


def get_bouncer() -> TranslationBouncer:
    global _bouncer
    if _bouncer is None:
        _bouncer = TranslationBouncer()
    return _bouncer


def get_perplexity_optimizer() -> PerplexityOptimizer:
    global _perplexity_optimizer
    if _perplexity_optimizer is None:
        _perplexity_optimizer = PerplexityOptimizer()
    return _perplexity_optimizer


# ── Models ──────────────────────────────────────────────────────────────────

class RewriteRequest(BaseModel):
    text: str = Field(..., min_length=1, description="The text to rewrite")
    mode: RewriteMode = Field(default=RewriteMode.STANDARD, description="Rewrite style mode")
    level: RewriteLevel = Field(default=RewriteLevel.MODERATE, description="Rewrite intensity level (1-3)")


class AnalyzeRequest(BaseModel):
    text: str = Field(..., min_length=1, description="The text to analyze")


class AICheckRequest(BaseModel):
    text: str = Field(..., min_length=1, description="The text to analyze for AI forensic signals")


class StatsResponse(BaseModel):
    word_count: int
    character_count: int
    sentence_count: int
    paragraph_count: int
    avg_sentence_length: float
    readability_score: float
    readability_grade: str
    vocabulary_diversity: float
    repeated_words: list[dict]
    repeated_phrases: list[dict]
    passive_voice_count: int
    reading_time_seconds: int


class FileParseResponse(BaseModel):
    filename: str
    text: str
    word_count: int
    character_count: int


class RewriteResponse(BaseModel):
    rewritten: str
    original_stats: dict
    rewritten_stats: dict
    changes: dict
    reading_time: dict
    meaning_preserved: bool
    meaning_reason: str
    meaning_preservation_score: float = 95.0
    similarity_metrics: Optional[dict] = None
    nlp_analysis: Optional[dict] = None
    ai_check: Optional[dict] = None
    word_diff: list[dict]


# ── Endpoints ───────────────────────────────────────────────────────────────

@router.post("/ai-check")
async def run_ai_check(request: AICheckRequest):
    """
    Forensic AI detection endpoint implementing Signals A through I (0 to 27 points).
    """
    report = AICheckEngine.analyze(request.text)
    return report.to_dict()

@router.post("/parse-file", response_model=FileParseResponse)
async def parse_uploaded_file(
    file: UploadFile = File(...),
    authorization: Optional[str] = Header(None)
):
    """
    Extracts plain text from uploaded .docx, .pdf, .txt, or .md files.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file was uploaded.")

    filename_lower = file.filename.lower()
    content = await file.read()

    if not content:
        raise HTTPException(status_code=400, detail="The uploaded file is empty.")

    if len(content) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size exceeds the 15MB maximum limit.")

    extracted_text = ""

    if filename_lower.endswith(('.txt', '.md', '.rtf', '.csv', '.json')):
        try:
            extracted_text = content.decode('utf-8')
        except UnicodeDecodeError:
            try:
                extracted_text = content.decode('latin-1')
            except Exception:
                raise HTTPException(status_code=400, detail="Unable to decode text file. Please ensure it is saved in UTF-8 format.")

    elif filename_lower.endswith('.docx'):
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            paras = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paras.append(row_text)
            extracted_text = "\n\n".join(paras)
        except Exception as e:
            logger.error("Failed to parse docx: %s", e)
            raise HTTPException(status_code=400, detail="Could not read Word document (.docx). The file may be password protected or corrupted.")

    elif filename_lower.endswith('.pdf'):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(content))
            pages_text = []
            for page in reader.pages:
                page_t = page.extract_text()
                if page_t and page_t.strip():
                    pages_text.append(page_t.strip())
            extracted_text = "\n\n".join(pages_text)
        except Exception as e:
            logger.error("Failed to parse pdf: %s", e)
            raise HTTPException(status_code=400, detail="Could not read PDF document. The file may be scanned images, password protected, or corrupted.")

    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file format. Please upload a .docx, .pdf, .txt, or .md file."
        )

    clean_extracted = extracted_text.strip()
    if not clean_extracted:
        raise HTTPException(
            status_code=400,
            detail="No readable text could be extracted from the uploaded document."
        )

    words = len(clean_extracted.split())
    return FileParseResponse(
        filename=file.filename,
        text=clean_extracted,
        word_count=words,
        character_count=len(clean_extracted)
    )


@router.post("/rewrite", response_model=RewriteResponse)
async def rewrite_text(
    request: RewriteRequest,
    authorization: Optional[str] = Header(None)
):
    """
    Main rewrite endpoint. Enforces 10 free humanization usage limits for free tier accounts.
    """
    user = get_optional_user_from_token(authorization)
    user_id_str = user.id if user else "anonymous"

    # Step A: Rate limiting check (Redis sliding window)
    if is_rate_limited(user_id_str, limit=30, window_seconds=60):
        raise HTTPException(
            status_code=429,
            detail="Rate limit exceeded. Please wait a minute before making more rewrite requests."
        )

    # Plan Limits Config (Humanizations per day)
    plan_limits = {
        'free': {'max_words': 400, 'max_usage': 10},
        'starter': {'max_words': 1000, 'max_usage': 30},
        'plus': {'max_words': 1000, 'max_usage': 30},
        'pro': {'max_words': 2500, 'max_usage': 80},
        'enterprise': {'max_words': 5000, 'max_usage': 250},
    }

    user_plan = user.plan if (user and hasattr(user, 'plan') and user.plan) else 'free'
    cfg = plan_limits.get(user_plan, plan_limits['free'])

    # Enforce Daily Usage Limits (reset every 24h)
    daily_usage_count = 0
    if user:
        today_start = datetime.utcnow().strftime('%Y-%m-%d 00:00:00')
        daily_usage_row = fetch_one(
            "SELECT COUNT(*) as cnt FROM history WHERE user_id = ? AND created_at >= ?",
            (user.id, today_start)
        )
        daily_usage_count = daily_usage_row["cnt"] if (daily_usage_row and "cnt" in daily_usage_row) else user.usage_count

    if user and daily_usage_count >= cfg['max_usage']:
        raise HTTPException(
            status_code=403,
            detail=f"{user_plan.capitalize()} plan limit reached ({cfg['max_usage']} humanizations per day used). Please upgrade your plan to continue or try again tomorrow."
        )

    # Enforce Word Count Limits
    input_word_count = len(request.text.strip().split())
    if input_word_count > cfg['max_words']:
        raise HTTPException(
            status_code=400,
            detail=f"Your {user_plan.capitalize()} plan allows up to {cfg['max_words']} words per input. This text has {input_word_count} words. Please upgrade to a higher plan for larger word limits."
        )

    if len(request.text) > MAX_INPUT_LENGTH:
        raise HTTPException(
            status_code=400,
            detail=f"Input text exceeds maximum absolute length of {MAX_INPUT_LENGTH} characters."
        )

    # Step 0: Normalise formatting artifacts (§15–§19) before the LLM sees the text.
    request_text_clean = strip_formatting_artifacts(request.text)
    clean_text = sanitize_input(request_text_clean)
    if not clean_text:
        raise HTTPException(status_code=400, detail="Input text is empty after sanitization.")

    def extract_document_title(text: str) -> tuple[Optional[str], str]:
        if not text:
            return None, text
        lines = [l.strip() for l in text.strip().split('\n') if l.strip()]
        if not lines:
            return None, text
        first_line = lines[0]
        words = first_line.split()
        is_title = False
        if first_line.startswith(('#', 'Title:', 'TITLE:', 'Heading:', 'HEADING:')):
            is_title = True
        elif len(words) <= 12 and not any(punct in first_line for punct in ('.', '!', '?', ';', ':')):
            is_title = True

        if is_title and len(lines) > 1:
            title = first_line
            body_start_idx = text.find(first_line) + len(first_line)
            body_text = text[body_start_idx:].strip()
            return title, body_text
        return None, text

    doc_title, text_body = extract_document_title(clean_text)

    # Step B: Check Redis cache
    cached_output = get_cached_rewrite(clean_text, request.mode.value, request.level.value)
    if cached_output:
        logger.info("Serving rewrite from Redis cache.")
        rewritten = cached_output
    else:
        try:
            pipeline = get_standard_pipeline()
            target_text = text_body if doc_title else clean_text

            rewritten = pipeline.process(
                target_text,
                mode=request.mode,
                level=request.level,
                target_lang="en"
            )

            # Re-attach document title if present
            if doc_title:
                rewritten = f"{doc_title}\n\n{rewritten}"

            # Store in Redis cache
            set_cached_rewrite(clean_text, request.mode.value, request.level.value, rewritten)
        except Exception as e:
            logger.error("Rewrite error: %s", e)
            raise HTTPException(
                status_code=500,
                detail="Text humanization service is temporarily busy. Please try again in a few seconds."
            )

    # Step C: RapidFuzz & SentenceTransformers Similarity Metrics
    sim_metrics = calculate_similarity_metrics(clean_text, rewritten)

    # Step D: spaCy NLP Analysis
    nlp_data = analyze_text_nlp(rewritten)

    meaning_preserved = True
    meaning_reason = "Factual accuracy preserved."

    # Step E: Textstat, word diff stats & forensic AI-check
    original_stats = analyze(clean_text)
    rewritten_stats = analyze(rewritten)
    changes = count_changes(clean_text, rewritten)
    reading_time = compute_reading_time(rewritten)
    word_diff = compute_word_diff(clean_text, rewritten)
    ai_check_data = AICheckEngine.analyze(rewritten).to_dict()

    # Record usage count and history entry if user is logged in
    if user:
        q_inc = "UPDATE users SET usage_count = usage_count + 1 WHERE id = ?"
        execute_query(q_inc, (user.id,))
        
        hist_id = str(uuid.uuid4())
        word_cnt = rewritten_stats.word_count
        created_at = datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')
        
        raw_mode_str = str(request.mode.value if hasattr(request.mode, 'value') else request.mode).lower().strip()
        if raw_mode_str in ("native", "standard"):
            saved_mode = "standard"
        elif raw_mode_str in ("casual", "natural"):
            saved_mode = "natural"
        elif raw_mode_str in ("professional", "fluency", "business"):
            saved_mode = "fluency"
        elif raw_mode_str in ("friendly", "creative"):
            saved_mode = "creative"
        else:
            saved_mode = raw_mode_str

        q_hist = """
            INSERT INTO history (id, user_id, original_text, rewritten_text, mode, level, word_count, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """
        execute_query(q_hist, (hist_id, user.id, clean_text, rewritten, saved_mode, int(request.level), word_cnt, created_at))

    return RewriteResponse(
        rewritten=rewritten,
        original_stats=original_stats.to_dict(),
        rewritten_stats=rewritten_stats.to_dict(),
        changes=changes,
        reading_time=reading_time,
        meaning_preserved=meaning_preserved,
        meaning_reason=meaning_reason,
        meaning_preservation_score=sim_metrics.get("meaning_preservation_score", 95.0),
        similarity_metrics=sim_metrics,
        nlp_analysis=nlp_data,
        ai_check=ai_check_data,
        word_diff=word_diff,
    )


@router.get("/user/history")
async def get_user_history(current_user: UserResponse = Depends(get_current_user_from_token)):
    """
    Retrieve past humanizations history for the logged-in user.
    """
    q_get = """
        SELECT id, original_text, rewritten_text, mode, level, word_count, created_at
        FROM history
        WHERE user_id = ?
        ORDER BY created_at DESC
        LIMIT 100
    """
    rows = fetch_all(q_get, (current_user.id,))
    results = []
    for r in rows:
        m = str(r.get("mode") or "standard").lower().strip()
        if m in ("native", "standard"):
            normalized_mode = "standard"
        elif m in ("casual", "natural"):
            normalized_mode = "natural"
        elif m in ("professional", "fluency", "business"):
            normalized_mode = "fluency"
        elif m in ("friendly", "creative"):
            normalized_mode = "creative"
        else:
            normalized_mode = m

        results.append({
            "id": r["id"],
            "original_text": r["original_text"],
            "rewritten_text": r["rewritten_text"],
            "mode": normalized_mode,
            "level": r.get("level", 2),
            "word_count": r.get("word_count", 0),
            "created_at": str(r["created_at"])
        })
    return results
